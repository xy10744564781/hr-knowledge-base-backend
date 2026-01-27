import tempfile
import os
import mimetypes
from typing import List, Tuple, Optional
from fastapi import UploadFile, HTTPException
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader, UnstructuredPDFLoader
from langchain_chroma import Chroma
from langchain.schema import Document
from logging_setup import logger
from config import (
    MAX_FILE_SIZE, SUPPORTED_FORMATS, CHROMA_DB_PATH, CHROMA_COLLECTION_NAME,
    DASHSCOPE_API_KEY, DASHSCOPE_BASE_URL, EMBEDDING_MODEL
)
from document_splitter import create_hr_splitter
from dashscope_embeddings import DashScopeEmbeddings


class ExcelLoader:
    """Excel 文件加载器 - 优化版"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载 Excel 文件并转换为 Document 对象"""
        try:
            import pandas as pd
            import numpy as np
            
            # 读取 Excel 文件的所有 sheet
            excel_file = pd.ExcelFile(self.file_path)
            documents = []
            
            for sheet_name in excel_file.sheet_names:
                # 读取每个 sheet
                df = pd.read_excel(self.file_path, sheet_name=sheet_name)
                
                # 跳过空 sheet
                if df.empty:
                    continue
                
                # 构建结构化内容
                content_parts = []
                
                # === 表格开始标记 ===
                content_parts.append(f"=== Excel 表格：{sheet_name} ===\n\n")
                
                # 列名信息
                columns = [str(col) for col in df.columns]
                content_parts.append(f"【列名】\n{', '.join(columns)}\n\n")
                
                # 数据统计信息（方案2）
                content_parts.append(f"【数据概览】\n")
                content_parts.append(f"总行数: {len(df)}\n")
                content_parts.append(f"总列数: {len(df.columns)}\n")
                
                # 为数值列添加统计信息
                numeric_cols = df.select_dtypes(include=[np.number]).columns
                if len(numeric_cols) > 0:
                    content_parts.append(f"\n数值列统计:\n")
                    for col in numeric_cols:
                        if df[col].notna().sum() > 0:  # 有非空值
                            content_parts.append(
                                f"  • {col}: "
                                f"最小值={df[col].min():.2f}, "
                                f"最大值={df[col].max():.2f}, "
                                f"平均值={df[col].mean():.2f}\n"
                            )
                
                # 为分类列添加分布信息
                categorical_cols = df.select_dtypes(include=['object']).columns
                if len(categorical_cols) > 0:
                    content_parts.append(f"\n分类列分布:\n")
                    for col in categorical_cols:
                        if df[col].notna().sum() > 0:
                            value_counts = df[col].value_counts()
                            if len(value_counts) <= 10:  # 只显示不超过10个类别
                                dist = ", ".join([f"{k}({v})" for k, v in value_counts.items()])
                                content_parts.append(f"  • {col}: {dist}\n")
                
                content_parts.append(f"\n【详细数据】\n")
                
                # 逐行添加数据（方案1优化）
                for idx, row in df.iterrows():
                    row_text = " | ".join(
                        f"{col}: {val}" 
                        for col, val in row.items() 
                        if pd.notna(val)
                    )
                    if row_text:  # 只添加非空行
                        content_parts.append(f"第{idx+1}行: {row_text}\n")
                
                # === 表格结束标记 ===
                content_parts.append(f"\n=== 表格结束 ===\n")
                
                content = "".join(content_parts)
                
                # 创建 Document 对象，包含丰富的 metadata
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": self.file_path,
                        "sheet_name": sheet_name,
                        "row_count": len(df),
                        "column_count": len(df.columns),
                        "columns": columns,
                        "has_numeric_data": len(numeric_cols) > 0,
                        "numeric_columns": list(numeric_cols) if len(numeric_cols) > 0 else []
                    }
                )
                documents.append(doc)
            
            if not documents:
                # 如果没有有效的 sheet，创建一个空文档
                logger.warning(f"Excel 文件 {self.file_path} 没有有效数据")
                documents.append(Document(
                    page_content="此 Excel 文件为空或无法读取有效数据",
                    metadata={"source": self.file_path}
                ))
            
            return documents
            
        except ImportError:
            raise ImportError("需要安装 pandas 和 openpyxl 库来处理 Excel 文件。请运行: pip install pandas openpyxl")
        except Exception as e:
            logger.error(f"加载 Excel 文件失败: {str(e)}")
            raise ValueError(f"Excel 文件处理失败: {str(e)}")


class DocxTableLoader:
    """Word 文档加载器 - 支持表格提取"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载 Word 文档，提取文本和表格"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(self.file_path)
            documents = []
            
            # 用于存储文档内容的各个部分
            content_parts = []
            table_count = 0
            
            # 遍历文档的所有元素（段落和表格）
            for element in doc.element.body:
                # 处理段落
                if element.tag.endswith('p'):
                    # 找到对应的段落对象
                    for para in doc.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text:
                                content_parts.append(text + "\n\n")
                            break
                
                # 处理表格
                elif element.tag.endswith('tbl'):
                    # 找到对应的表格对象
                    for table in doc.tables:
                        if table._element == element:
                            table_count += 1
                            table_text = self._extract_table(table, table_count)
                            if table_text:
                                content_parts.append(table_text)
                            break
            
            # 如果没有内容，使用备用方法（纯文本提取）
            if not content_parts:
                logger.warning(f"Word文档 {self.file_path} 使用备用方法提取")
                all_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if all_text:
                    content_parts.append(all_text)
            
            # 合并所有内容
            full_content = "".join(content_parts)
            
            if not full_content.strip():
                logger.warning(f"Word文档 {self.file_path} 没有有效内容")
                full_content = "此 Word 文档为空或无法读取有效内容"
            
            # 创建 Document 对象
            doc_obj = Document(
                page_content=full_content,
                metadata={
                    "source": self.file_path,
                    "table_count": table_count,
                    "has_tables": table_count > 0
                }
            )
            documents.append(doc_obj)
            
            return documents
            
        except ImportError:
            raise ImportError("需要安装 python-docx 库来处理 Word 文件。请运行: pip install python-docx")
        except Exception as e:
            logger.error(f"加载 Word 文件失败: {str(e)}")
            raise ValueError(f"Word 文件处理失败: {str(e)}")
    
    def _extract_table(self, table, table_number: int) -> str:
        """提取表格内容并格式化"""
        try:
            content_parts = []
            
            # === 表格开始标记 ===
            content_parts.append(f"\n=== Word 表格 {table_number} ===\n\n")
            
            # 提取表头（第一行）
            if len(table.rows) > 0:
                headers = []
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        headers.append(header_text)
                
                if headers:
                    content_parts.append(f"【列名】\n{' | '.join(headers)}\n\n")
                    content_parts.append(f"【表格数据】\n")
                    
                    # 提取数据行（从第二行开始）
                    for row_idx, row in enumerate(table.rows[1:], start=1):
                        row_data = []
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                # 如果有表头，使用表头名称
                                if cell_idx < len(headers):
                                    row_data.append(f"{headers[cell_idx]}: {cell_text}")
                                else:
                                    row_data.append(cell_text)
                        
                        if row_data:
                            content_parts.append(f"第{row_idx}行: {' | '.join(row_data)}\n")
                else:
                    # 没有明确表头，直接输出所有行
                    content_parts.append(f"【表格数据】\n")
                    for row_idx, row in enumerate(table.rows, start=1):
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            content_parts.append(f"第{row_idx}行: {' | '.join(row_data)}\n")
            
            # === 表格结束标记 ===
            content_parts.append(f"\n=== 表格结束 ===\n\n")
            
            return "".join(content_parts)
            
        except Exception as e:
            logger.error(f"提取表格失败: {str(e)}")
            return f"\n[表格 {table_number} 提取失败]\n\n"


# 全局向量存储实例
vector_store = None

def init_chroma():
    """初始化ChromaDB向量数据库（使用LangChain + 阿里云API）"""
    global vector_store
    if vector_store is None:
        try:
            # 使用自定义的DashScopeEmbeddings适配器
            embeddings = DashScopeEmbeddings(
                api_key=DASHSCOPE_API_KEY,
                model=EMBEDDING_MODEL,
                base_url=DASHSCOPE_BASE_URL
            )
            
            logger.info(f"配置阿里云Embedding模型: {EMBEDDING_MODEL}")
            logger.info(f"API Base URL: {DASHSCOPE_BASE_URL}")
            
            # 创建或加载ChromaDB向量存储
            vector_store = Chroma(
                collection_name=CHROMA_COLLECTION_NAME,
                embedding_function=embeddings,
                persist_directory=CHROMA_DB_PATH
            )
            
            logger.info(f"ChromaDB初始化成功（集合: {CHROMA_COLLECTION_NAME}）")
            
        except Exception as e:
            logger.error(f"ChromaDB初始化失败: {e}", exc_info=True)
            vector_store = None
            
    return vector_store

class VectorManager:
    """向量管理器包装类（基于LangChain Chroma）"""
    
    def __init__(self, vector_store):
        self.vector_store = vector_store
        self.document_cache = []  # 缓存所有文档用于BM25检索（dev-mix新增）
        self._load_existing_documents()  # 加载已有文档到缓存
    
    def _load_existing_documents(self):
        """加载已有文档到缓存（用于BM25）"""
        try:
            if not self.vector_store:
                return
            
            results = self.vector_store.get()
            if results and results.get('documents') and results.get('metadatas'):
                from langchain.schema import Document
                for doc_text, metadata in zip(results['documents'], results['metadatas']):
                    self.document_cache.append(
                        Document(page_content=doc_text, metadata=metadata or {})
                    )
                logger.info(f"已加载 {len(self.document_cache)} 个文档到缓存（用于BM25）")
        except Exception as e:
            logger.warning(f"加载文档缓存失败: {e}")
    
    def get_all_documents(self):
        """获取所有缓存的文档（用于BM25检索）"""
        return self.document_cache
    
    def check_duplicate_title(self, title, exclude_doc_id=None):
        """检查标题是否重复"""
        if not self.vector_store:
            return False
            
        try:
            # 使用LangChain的get方法查询
            results = self.vector_store.get()
            
            if not results or not results.get('metadatas'):
                return False
                
            # 检查是否有相同标题的文档
            for i, metadata in enumerate(results['metadatas']):
                if metadata and metadata.get('title') == title:
                    # 如果指定了排除的文档ID，跳过该文档
                    if exclude_doc_id and results.get('ids') and results['ids'][i].startswith(exclude_doc_id):
                        continue
                    return True
                    
            return False
            
        except Exception as e:
            logger.error(f"检查重复标题时出错: {e}")
            return False
    
    def add_document(self, texts, metadatas, document_id=None, ids=None):
        """添加文档"""
        if not self.vector_store:
            raise Exception("向量存储未初始化")
        
        try:
            from langchain.schema import Document
            
            # 确保texts是字符串列表
            if not isinstance(texts, list):
                texts = [texts]
            
            # 确保metadatas是字典列表
            if not isinstance(metadatas, list):
                metadatas = [metadatas]
            
            # 过滤空文本并确保格式正确
            valid_texts = []
            valid_metadatas = []
            for i, text in enumerate(texts):
                if text and isinstance(text, str) and text.strip():
                    # 确保文本是纯字符串，移除特殊字符
                    clean_text = str(text).strip()
                    # 限制单个文本长度，避免API限制
                    if len(clean_text) > 8000:
                        logger.warning(f"文本块 {i} 过长({len(clean_text)}字符)，截断到8000字符")
                        clean_text = clean_text[:8000]
                    
                    valid_texts.append(clean_text)
                    valid_metadatas.append(metadatas[i] if i < len(metadatas) else {})
            
            if not valid_texts:
                logger.warning("没有有效的文本内容可添加")
                return False
            
            logger.info(f"准备添加 {len(valid_texts)} 个有效文本块")
            
            # 生成IDs
            if ids is None:
                if document_id:
                    ids = [f"{document_id}_chunk_{i}" for i in range(len(valid_texts))]
                else:
                    import uuid
                    base_id = str(uuid.uuid4())
                    ids = [f"{base_id}_chunk_{i}" for i in range(len(valid_texts))]
            
            # 分批添加，避免单次请求过大
            batch_size = 10
            for i in range(0, len(valid_texts), batch_size):
                batch_texts = valid_texts[i:i+batch_size]
                batch_metadatas = valid_metadatas[i:i+batch_size]
                batch_ids = ids[i:i+batch_size]
                
                logger.info(f"添加批次 {i//batch_size + 1}/{(len(valid_texts)-1)//batch_size + 1}，包含 {len(batch_texts)} 个文本")
                
                # 使用add_texts方法
                self.vector_store.add_texts(
                    texts=batch_texts,
                    metadatas=batch_metadatas,
                    ids=batch_ids
                )
            
            # 同时添加到文档缓存（用于BM25）- dev-mix新增
            docs = [Document(page_content=t, metadata=m) 
                   for t, m in zip(valid_texts, valid_metadatas)]
            self.document_cache.extend(docs)
            logger.info(f"已添加 {len(docs)} 个文档到缓存")
            
            logger.info(f"成功添加 {len(valid_texts)} 个文档到向量存储")
            return True
            
        except Exception as e:
            logger.error(f"添加文档时出错: {e}", exc_info=True)
            return False
    
    def delete_document(self, document_id):
        """删除文档"""
        if not self.vector_store:
            raise Exception("向量存储未初始化")
        
        try:
            # 查找所有相关的chunk IDs
            results = self.vector_store.get()
            if results and results.get('ids'):
                chunk_ids = [
                    id for id in results['ids']
                    if id.startswith(document_id)
                ]
                
                if chunk_ids:
                    self.vector_store.delete(ids=chunk_ids)
                    logger.info(f"删除文档 {document_id}，共 {len(chunk_ids)} 个chunks")
                    
                    # 同时从缓存中删除（dev-mix新增）
                    self.document_cache = [
                        doc for doc in self.document_cache
                        if not (hasattr(doc, 'metadata') and 
                               doc.metadata.get('document_id') == document_id)
                    ]
                    logger.info(f"已从缓存中删除文档 {document_id}")
                    
                    return True
            return False
            
        except Exception as e:
            logger.error(f"删除文档时出错: {e}")
            return False
    
    def update_document(self, document_id, texts, metadatas, ids=None):
        """更新文档"""
        if not self.vector_store:
            raise Exception("向量存储未初始化")
        
        try:
            # 先删除旧的chunks
            self.delete_document(document_id)
            # 再添加新的chunks
            return self.add_document(texts, metadatas, document_id, ids)
        except Exception as e:
            logger.error(f"更新文档时出错: {e}")
            return False
    
    def search_documents(self, query, k=10, filter_metadata=None):
        """搜索文档（使用LangChain接口）"""
        if not self.vector_store:
            raise Exception("向量存储未初始化")
        
        try:
            # 使用LangChain的similarity_search_with_score
            results = self.vector_store.similarity_search_with_score(
                query,
                k=k,
                filter=filter_metadata
            )
            
            # 转换结果格式，添加分数到metadata
            documents = []
            for doc, score in results:
                # 将距离转换为相似度分数
                similarity_score = max(0.0, 1.0 - (score / 2.0))
                
                if not hasattr(doc, 'metadata'):
                    doc.metadata = {}
                doc.metadata['score'] = similarity_score
                
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            logger.error(f"搜索文档时出错: {e}", exc_info=True)
            return []
    
    def list_documents(self, limit=None):
        """列出文档"""
        if not self.vector_store:
            return []
        
        try:
            results = self.vector_store.get()
            if not results or not results.get('metadatas'):
                return []
            
            # 按document_id分组，避免重复
            documents = {}
            for metadata in results['metadatas']:
                if metadata and 'document_id' in metadata:
                    doc_id = metadata['document_id']
                    if doc_id not in documents:
                        documents[doc_id] = metadata
            
            doc_list = list(documents.values())
            if limit:
                doc_list = doc_list[:limit]
                
            return doc_list
            
        except Exception as e:
            logger.error(f"列出文档时出错: {e}")
            return []
    
    def get_collection_stats(self):
        """获取集合统计信息"""
        if not self.vector_store:
            return {"total_documents": 0}
        try:
            results = self.vector_store.get()
            count = len(results.get('ids', [])) if results else 0
            return {
                "total_documents": count,
                "collection_name": CHROMA_COLLECTION_NAME
            }
        except Exception as e:
            logger.error(f"获取统计信息时出错: {e}")
            return {"total_documents": 0}

def get_vector_manager():
    """获取向量管理器"""
    store = init_chroma()
    return VectorManager(store) if store else None

def validate_file(file: UploadFile) -> Tuple[bool, str]:
    """
    验证上传文件的有效性
    返回: (是否有效, 错误信息)
    """
    try:
        # 检查文件名
        if not file.filename:
            return False, "文件名不能为空"
        
        # 检查文件扩展名
        file_ext = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        if not file_ext:
            return False, "文件必须有扩展名"
        
        if f".{file_ext}" not in SUPPORTED_FORMATS:
            return False, f"不支持的文件格式: .{file_ext}。支持的格式: {', '.join(SUPPORTED_FORMATS)}"
        
        # 检查文件大小
        file.file.seek(0, 2)  # 移动到文件末尾
        file_size = file.file.tell()
        file.file.seek(0)  # 重置到文件开头
        
        if file_size == 0:
            return False, "文件不能为空"
        
        if file_size > MAX_FILE_SIZE:
            return False, f"文件大小超过限制 ({file_size / (1024*1024):.1f}MB > {MAX_FILE_SIZE / (1024*1024)}MB)"
        
        # 检查MIME类型（如果可能）
        mime_type, _ = mimetypes.guess_type(file.filename)
        if mime_type:
            allowed_mimes = {
                'application/pdf': ['.pdf'],
                'text/plain': ['.txt', '.md', '.log'],
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
                'application/msword': ['.doc']
            }
            
            valid_mime = False
            for allowed_mime, extensions in allowed_mimes.items():
                if mime_type == allowed_mime and f".{file_ext}" in extensions:
                    valid_mime = True
                    break
            
            if not valid_mime and mime_type not in ['text/x-python', 'application/javascript', 'application/json']:
                logger.warning(f"MIME类型可能不匹配: {mime_type} for {file.filename}")
        
        return True, ""
        
    except Exception as e:
        logger.error(f"文件验证失败: {str(e)}")
        return False, f"文件验证失败: {str(e)}"

def get_optimal_chunk_settings(file_ext: str, file_size: int) -> Tuple[int, int]:
    """
    根据文件类型和大小获取最优的分块设置
    返回: (chunk_size, chunk_overlap)
    """
    # 基础设置
    base_chunk_size = 1000
    base_overlap = 200
    
    # 根据文件类型调整
    if file_ext == 'pdf':
        # PDF通常内容密度较高
        chunk_size = 1200
        overlap = 250
    elif file_ext in ['docx', 'doc']:
        # Word文档通常格式化较好
        chunk_size = 1000
        overlap = 200
    elif file_ext in ['md', 'txt']:
        # 纯文本文档
        chunk_size = 800
        overlap = 150
    else:
        # 代码文件等
        chunk_size = 600
        overlap = 100
    
    # 根据文件大小调整
    if file_size > 5 * 1024 * 1024:  # 大于5MB
        chunk_size = int(chunk_size * 1.2)
        overlap = int(overlap * 1.1)
    elif file_size < 100 * 1024:  # 小于100KB
        chunk_size = int(chunk_size * 0.8)
        overlap = int(overlap * 0.8)
    
    return chunk_size, overlap

def create_document_loader(temp_path: str, file_ext: str, filename: str):
    """
    根据文件类型创建合适的文档加载器
    """
    try:
        if file_ext == 'pdf':
            # 优先使用PyPDFLoader
            return PyPDFLoader(temp_path)
        elif file_ext in ["txt", "md", "c", "h", "cpp", "hpp", "java", "js", "ts", "jsx", "tsx",
                          "vue", "py", "go", "rs", "sql", "json", "yml", "yaml", "cfg", "ini", "log"]:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']
            for encoding in encodings:
                try:
                    return TextLoader(temp_path, encoding=encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"无法使用任何编码读取文件: {filename}")
        elif file_ext in ['docx', 'doc']:
            # 使用支持表格提取的 DocxTableLoader
            return DocxTableLoader(temp_path)
        elif file_ext in ['xlsx', 'xls']:
            # Excel 文件处理
            return ExcelLoader(temp_path)
        else:
            raise ValueError(f'不支持的文档格式: {file_ext}')
    except Exception as e:
        logger.error(f"创建文档加载器失败: {str(e)}")
        raise

def process_upload_file(file: UploadFile) -> List[str]:
    """
    处理上传的文件并提取文本内容
    使用智能语义切割器
    """
    logger.info(f"开始处理文件: {file.filename}")
    
    # 1. 文件验证
    is_valid, error_msg = validate_file(file)
    if not is_valid:
        logger.error(f"文件验证失败: {error_msg}")
        raise ValueError(error_msg)
    
    # 2. 获取文件信息
    file_ext = file.filename.split('.')[-1].lower()
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    
    logger.info(f"文件信息: 类型={file_ext}, 大小={file_size / 1024:.1f}KB")
    
    # 3. 创建临时文件
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as temp_file:
            # 分块读取文件内容，避免内存问题
            chunk_size = 8192
            while True:
                chunk = file.file.read(chunk_size)
                if not chunk:
                    break
                temp_file.write(chunk)
            temp_path = temp_file.name
        
        logger.info(f"临时文件创建成功: {temp_path}")
        
        # 4. 创建文档加载器
        loader = create_document_loader(temp_path, file_ext, file.filename)
        
        # 5. 加载文档
        try:
            documents = loader.load()
            logger.info(f"文档加载成功，原始页数: {len(documents)}")
            
            if not documents:
                raise ValueError("文档内容为空或无法读取")
            
            # 检查文档内容
            total_content = sum(len(doc.page_content) for doc in documents)
            if total_content < 10:
                raise ValueError("文档内容过少，可能是空文档或格式不正确")
            
            # 6. 使用智能语义切割器
            from document_splitter import create_hr_splitter
            
            splitter = create_hr_splitter()
            
            # 记录原始文档内容长度
            for i, doc in enumerate(documents):
                logger.info(f"原始文档 {i+1} 内容长度: {len(doc.page_content)} 字符")
                logger.info(f"原始文档 {i+1} 前100字符: {doc.page_content[:100]}")
            
            # 调用split_documents
            logger.info("开始调用splitter.split_documents...")
            split_docs = splitter.split_documents(documents)
            logger.info(f"文档分割完成，分块数: {len(split_docs)}")
            
            # 记录每个分块的长度
            for i, doc in enumerate(split_docs[:10]):  # 记录前10个
                logger.info(f"分块 {i+1} 长度: {len(doc.page_content)} 字符，前50字符: {doc.page_content[:50]}")

            
            # 7. 提取文本内容
            valid_chunks = []
            for doc in split_docs:
                content = doc.page_content.strip()
                if len(content) >= 20:  # 至少20个字符
                    valid_chunks.append(content)
            
            if not valid_chunks:
                raise ValueError("文档分割后没有有效内容块")
            
            logger.info(f"有效内容块数量: {len(valid_chunks)}")
            return valid_chunks
            
        except Exception as e:
            logger.error(f'主要加载器失败: {str(e)}')
            
            # 8. 备用处理方法
            if file_ext == 'pdf':
                try:
                    logger.warning("尝试使用备用PDF处理方法")
                    loader = UnstructuredPDFLoader(temp_path)
                    documents = loader.load()
                    
                    if documents:
                        from document_splitter import create_hr_splitter
                        splitter = create_hr_splitter()
                        split_docs = splitter.split_documents(documents)
                        
                        valid_chunks = [doc.page_content.strip() for doc in split_docs 
                                      if len(doc.page_content.strip()) >= 20]
                        if valid_chunks:
                            logger.info(f"备用方法成功，有效块数: {len(valid_chunks)}")
                            return valid_chunks
                    
                except Exception as backup_e:
                    logger.error(f'备用PDF处理方法也失败: {str(backup_e)}')
            
            # 如果所有方法都失败，抛出详细错误
            raise ValueError(f'文件处理失败: {str(e)}。请检查文件格式是否正确，或尝试转换为其他支持的格式。')
    
    except ValueError:
        # 重新抛出验证错误
        raise
    except Exception as e:
        logger.error(f'文件处理过程中发生未预期错误: {str(e)}', exc_info=True)
        raise ValueError(f'文件处理失败: {str(e)}')
    
    finally:
        # 9. 清理临时文件
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
                logger.info("临时文件清理完成")
            except Exception as e:
                logger.warning(f"临时文件清理失败: {str(e)}")

def get_document_info(file: UploadFile) -> dict:
    """
    获取文档基本信息，不进行实际处理
    """
    try:
        file_ext = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        
        file.file.seek(0, 2)
        file_size = file.file.tell()
        file.file.seek(0)
        
        mime_type, _ = mimetypes.guess_type(file.filename)
        
        return {
            "filename": file.filename,
            "extension": file_ext,
            "size_bytes": file_size,
            "size_mb": round(file_size / (1024 * 1024), 2),
            "mime_type": mime_type,
            "is_supported": f".{file_ext}" in SUPPORTED_FORMATS,
            "estimated_chunks": max(1, file_size // 1000)  # 粗略估计
        }
    except Exception as e:
        logger.error(f"获取文档信息失败: {str(e)}")
        return {"error": str(e)}
